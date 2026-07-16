"""
CiteFlow AI – Flask Application
--------------------------------
Two modes: Citation Assistant & Proofreading Assistant
"""
import os
import re
import json
import time
import io
import difflib
import requests
from flask import Flask, session, request, jsonify, render_template, redirect, url_for, send_file
from openai import OpenAI
import PyPDF2
import docx
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')

# ---------------------------------------------------------------------------
# CONSTANTS
# ---------------------------------------------------------------------------

CROSSREF_ENDPOINT = "https://api.crossref.org/works"
ABSTRACT_SNIPPET_LENGTH = 250
QUOTE_SNIPPET_LENGTH = 150
MIN_PARAGRAPH_LENGTH = 100
KEYWORDS_PER_QUERY = 4

HEADING_PATTERNS = [
    r"^introduction\s*:?\s*$",
    r"^conclusion\s*:?\s*$",
    r"^abstract\s*:?\s*$",
    r"^methodology\s*:?\s*$",
    r"^results\s*:?\s*$",
    r"^discussion\s*:?\s*$",
    r"^references\s*:?\s*$",
    r"^background\s*:?\s*$",
]

DEFAULT_DOCUMENT_TEXT = """# Introduction

The convergence of operational technology and information technology has exposed industrial control systems to unprecedented cybersecurity risks. Modern cyber-physical systems increasingly adopt zero-trust architectures to mitigate lateral movement attacks, yet the integration of legacy SCADA protocols with cloud-based analytics introduces persistent vulnerabilities that adversaries can exploit through supply-chain compromise and firmware manipulation. Robust anomaly-detection frameworks combined with continuous authentication are now considered essential safeguards for critical infrastructure resilience.

# Autonomous Vehicles

Perception pipelines in autonomous vehicles rely heavily on multi-modal sensor fusion, combining LiDAR point clouds, radar returns, and high-resolution camera imagery to construct a coherent three-dimensional understanding of the driving environment. Recent advances in real-time deep-learning inference have enabled onboard computer-vision systems to perform object detection and trajectory prediction with sub-hundred-millisecond latency, although occlusion handling and adverse-weather robustness remain open engineering challenges for large-scale deployment.

# Biotechnology

Enzyme engineering through directed evolution has transformed industrial biocatalysis by allowing researchers to iteratively optimize protein stability, substrate specificity, and catalytic turnover without requiring complete a priori knowledge of the underlying protein-folding landscape. Coupling high-throughput screening with machine-learning-guided mutagenesis has accelerated the discovery of thermostable enzyme variants suitable for industrial-scale biomanufacturing under non-native reaction conditions.
"""

# ---------------------------------------------------------------------------
# HELPERS (shared)
# ---------------------------------------------------------------------------

def split_into_paragraphs(document_text: str) -> list:
    return re.split(r"\n\s*\n", document_text.strip("\n"))

def is_citable_paragraph(block_text: str) -> bool:
    stripped = block_text.strip()
    if stripped.startswith('#'):
        return False
    if len(stripped) < MIN_PARAGRAPH_LENGTH:
        return False
    for pattern in HEADING_PATTERNS:
        if re.match(pattern, stripped.lower()):
            return False
    return True

def extract_text_from_txt(file_bytes):
    return file_bytes.decode('utf-8', errors='ignore')

def extract_text_from_docx(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    lines = []
    
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
            
        style_name = p.style.name.lower()
        prefix = ""
        
        if style_name.startswith('heading 1'):
            prefix = "# "
        elif style_name.startswith('heading 2'):
            prefix = "## "
        elif style_name.startswith('heading 3'):
            prefix = "### "
        elif 'list bullet' in style_name:
            prefix = "- "
        elif 'list number' in style_name:
            prefix = "1. "
            
        para_md = ""
        for run in p.runs:
            text = run.text
            if not text.strip():
                para_md += text
                continue
                
            stripped_text = text.strip()
            if run.bold and run.italic:
                formatted_text = f"***{stripped_text}***"
            elif run.bold:
                formatted_text = f"**{stripped_text}**"
            elif run.italic:
                formatted_text = f"*{stripped_text}*"
            else:
                formatted_text = stripped_text
                
            if text.startswith(' '):
                formatted_text = ' ' + formatted_text
            if text.endswith(' '):
                formatted_text = formatted_text + ' '
                
            para_md += formatted_text
            
        lines.append(f"{prefix}{para_md.strip()}")
        
    return "\n\n".join(lines)

def extract_text_from_pdf(file_bytes):
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n\n"
    return text.strip()

def extract_text_from_upload(file_bytes, filename):
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext == 'txt':
        return extract_text_from_txt(file_bytes)
    elif ext == 'docx':
        return extract_text_from_docx(file_bytes)
    elif ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    else:
        raise ValueError("Unsupported file type. Please upload .txt, .docx, or .pdf.")

# ---------------------------------------------------------------------------
# MARKDOWN TO DOCX PARSER
# ---------------------------------------------------------------------------

def add_markdown_paragraph_to_docx(doc, md_text):
    stripped = md_text.strip()
    if not stripped:
        return
        
    para = None
    
    if stripped.startswith('### '):
        para = doc.add_heading(level=3)
        stripped = stripped[4:]
    elif stripped.startswith('## '):
        para = doc.add_heading(level=2)
        stripped = stripped[3:]
    elif stripped.startswith('# '):
        para = doc.add_heading(level=1)
        stripped = stripped[2:]
    elif stripped.startswith('- ') or stripped.startswith('* '):
        para = doc.add_paragraph(style='List Bullet')
        stripped = stripped[2:]
    elif re.match(r'^\d+\.\s', stripped):
        para = doc.add_paragraph(style='List Number')
        stripped = re.sub(r'^\d+\.\s', '', stripped)
    else:
        para = doc.add_paragraph()
        
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', stripped)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = para.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)

# ---------------------------------------------------------------------------
# AUTHOR FORMATTING (for citation mode)
# ---------------------------------------------------------------------------

def _get_surname_from_name(full_name: str) -> str:
    parts = full_name.strip().split()
    return parts[-1] if parts else "Unknown"

def format_authors_inline(paper: dict, style: str) -> str:
    authors = paper.get("authors") or []
    names = [a.get("name", "").strip() for a in authors if a.get("name")]
    if not names:
        return "Unknown"
    surnames = [_get_surname_from_name(name) for name in names]
    if style in ("APA", "Harvard"):
        if len(surnames) == 1:
            return surnames[0]
        elif len(surnames) == 2:
            return f"{surnames[0]} & {surnames[1]}"
        else:
            return f"{surnames[0]} et al."
    return surnames[0]

def get_all_author_names(paper: dict) -> str:
    authors = paper.get("authors") or []
    names = [a.get("name", "Unknown").strip() for a in authors if a.get("name")]
    if not names:
        return "Unknown Author"
    return ", ".join(names)

# ---------------------------------------------------------------------------
# CROSSREF INTEGRATION (cached)
# ---------------------------------------------------------------------------

_cached_papers = {}

def _fetch_uncached(search_query: str):
    params = {
        "query": search_query,
        "rows": 3,
        "select": "title,author,issued,container-title,abstract,URL"
    }
    headers = {"User-Agent": "CiteFlowAI/1.0 (mailto:contact@example.com)"}
    try:
        resp = requests.get(CROSSREF_ENDPOINT, params=params, headers=headers, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        items = data.get("message", {}).get("items", [])
        if not items:
            return None, "No articles found on Crossref for these keywords."
        
        papers = []
        for item in items:
            authors = []
            for a in item.get("author", []):
                if "family" in a and "given" in a:
                    name = f"{a['given']} {a['family']}".strip()
                elif "family" in a:
                    name = a['family'].strip()
                elif "given" in a:
                    name = a['given'].strip()
                elif "name" in a:
                    name = a['name'].strip()
                else:
                    continue
                if name:
                    authors.append({"name": name})
            year = "n.d."
            try:
                year = item.get("issued", {}).get("date-parts", [[None]])[0][0]
            except Exception:
                pass
                
            raw_abstract = item.get("abstract", "")
            clean_abstract = re.sub(r'<[^>]+>', '', raw_abstract) if raw_abstract else ""
            
            if len(clean_abstract) > 300:
                clean_abstract = clean_abstract[:300] + "..."
                
            paper = {
                "title": item.get("title", ["Untitled"])[0],
                "authors": authors,
                "venue": item.get("container-title", ["Unknown venue"])[0],
                "year": year,
                "abstract": clean_abstract,
                "url": item.get("URL", "")
            }
            papers.append(paper)
        return papers, None
    except requests.exceptions.RequestException as e:
        return None, f"Network error with Crossref: {e}"

def fetch_top_papers(search_query: str):
    now = time.time()
    if search_query in _cached_papers:
        entry = _cached_papers[search_query]
        if now - entry['timestamp'] < 600:
            return entry['papers'], entry['error']
    papers, error = _fetch_uncached(search_query)
    _cached_papers[search_query] = {
        'papers': papers,
        'error': error,
        'timestamp': now
    }
    return papers, error

# ---------------------------------------------------------------------------
# SCRAPING DE SECOURS (quand Crossref ne fournit pas d'abstract)
# ---------------------------------------------------------------------------

_cached_source_text = {}

def _scrape_paper_text(url: str) -> str:
    headers = {"User-Agent": "CiteFlowAI/1.0 (mailto:contact@example.com)"}
    resp = requests.get(url, headers=headers, timeout=8, allow_redirects=True)
    resp.raise_for_status()
    html = resp.text

    # 1. Métadonnées académiques standard (les plus fiables : vrai texte, pas de bruit de page)
    meta_patterns = [
        r'<meta[^>]+name=["\']citation_abstract["\'][^>]+content=["\']([^"\']+)["\']',
        r'<meta[^>]+name=["\']dc\.description["\'][^>]+content=["\']([^"\']+)["\']',
        r'<meta[^>]+property=["\']og:description["\'][^>]+content=["\']([^"\']+)["\']',
        r'<meta[^>]+name=["\']description["\'][^>]+content=["\']([^"\']+)["\']',
    ]
    for pattern in meta_patterns:
        match = re.search(pattern, html, re.IGNORECASE)
        if match:
            text = re.sub(r'<[^>]+>', '', match.group(1)).strip()
            text = re.sub(r'&amp;', '&', text)
            if len(text) > 50:
                return text[:1500]

    # 2. Repli : texte brut visible de la page (best-effort, bruité mais mieux que rien)
    text = re.sub(r'<script[^>]*>.*?</script>', ' ', html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', ' ', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text[:1500]

def get_citable_source_text(paper: dict) -> str:
    """
    Retourne le meilleur texte disponible pour extraire une vraie citation :
    1. l'abstract Crossref s'il existe,
    2. sinon un scraping de secours de la page du papier (URL, mis en cache),
    3. sinon chaîne vide.
    """
    abstract = paper.get("abstract", "").strip()
    if abstract:
        return abstract

    url = (paper.get("url") or "").strip()
    if not url:
        return ""

    now = time.time()
    cached = _cached_source_text.get(url)
    if cached and now - cached['timestamp'] < 600:
        return cached['text']

    try:
        text = _scrape_paper_text(url)
    except requests.exceptions.RequestException as e:
        app.logger.error(f"Scraping error for {url}: {e}")
        text = ""

    _cached_source_text[url] = {'text': text, 'timestamp': now}
    return text

# ---------------------------------------------------------------------------
# CITATION HELPERS
# ---------------------------------------------------------------------------

def build_inline_tag(paper: dict, style: str, number: int) -> str:
    if style == "IEEE":
        return f"[{number}]"
    year = paper.get("year") or "n.d."
    author_str = format_authors_inline(paper, style)
    return f"({author_str}, {year})"

def build_name_year_tag(paper: dict) -> str:
    """Toujours au format (Nom, Année), quel que soit le style de citation actif.
    Utilisé pour le mode Quote + Citation."""
    year = paper.get("year") or "n.d."
    author_str = format_authors_inline(paper, "APA")
    return f"({author_str}, {year})"

def build_inline_citation_text(paper: dict, style: str, mode: str, number: int) -> str:
    tag = build_inline_tag(paper, style, number)
    
    if mode == "Insert Quote + Citation":
        abstract = paper.get("abstract", "").strip()
        if not abstract:
            return f" {tag}"
            
        if len(abstract) > QUOTE_SNIPPET_LENGTH:
            snippet = abstract[:QUOTE_SNIPPET_LENGTH].rstrip() + "..."
        else:
            snippet = abstract
            
        return f' "{snippet}" {tag}'
    else:
        return f" {tag}"

def _normalize_for_match(s: str) -> str:
    return re.sub(r'\s+', ' ', s).strip().lower()

def _split_sentences(text: str) -> list:
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in sentences if s.strip()]

def _closest_faithful_sentence(candidate: str, abstract: str) -> str:
    """
    Garantit que la phrase retournée est une PHRASE RÉELLE et COMPLÈTE
    de l'abstract. Si le modèle a reformulé ou tronqué, on retombe sur
    la phrase de l'abstract la plus proche du candidat (jamais de
    citation inventée ou partielle).
    """
    sentences = _split_sentences(abstract)
    if not sentences:
        return candidate.strip().strip('"')

    norm_candidate = _normalize_for_match(candidate)

    # Correspondance exacte (le candidat EST une phrase de l'abstract)
    for sentence in sentences:
        if _normalize_for_match(sentence) == norm_candidate:
            return sentence

    # Le candidat est contenu dans une phrase (ou vice versa) -> on garde la phrase complète
    for sentence in sentences:
        norm_sentence = _normalize_for_match(sentence)
        if norm_candidate and (norm_candidate in norm_sentence or norm_sentence in norm_candidate):
            return sentence

    # Sinon, phrase la plus proche par similarité textuelle
    best_sentence = max(
        sentences,
        key=lambda s: difflib.SequenceMatcher(None, _normalize_for_match(s), norm_candidate).ratio()
    )
    return best_sentence

def extract_faithful_quote(paragraph: str, abstract: str) -> str:
    """
    Un seul appel API, sortie JSON structurée (parsing fiable, pas de
    guillemets à deviner), puis validation locale garantissant que la
    citation retournée est une PHRASE COMPLÈTE et réellement présente
    dans l'abstract (jamais un fragment inventé ou paraphrasé).
    """
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    prompt = (
        "Paragraph I'm writing:\n"
        f"\"{paragraph}\"\n\n"
        "Paper abstract:\n"
        f"\"{abstract}\"\n\n"
        "Task: pick the ONE full sentence, copied VERBATIM and IN FULL from "
        "the abstract (same words, same order, no edits, no truncation), that "
        "best supports my paragraph. Do not summarize, reword, or return a "
        "partial sentence.\n"
        "Respond only as JSON: {\"quote\": \"<one full sentence copied exactly from the abstract>\"}"
    )
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a precise academic assistant. You only ever copy full sentences verbatim, never paraphrase or truncate."},
            {"role": "user", "content": prompt}
        ],
        temperature=0,
        max_tokens=150,
        response_format={"type": "json_object"}
    )
    raw = response.choices[0].message.content.strip()
    try:
        parsed = json.loads(raw)
        candidate = str(parsed.get("quote", "")).strip()
    except (json.JSONDecodeError, AttributeError):
        candidate = raw.strip().strip('"')

    if not candidate:
        candidate = abstract

    return _closest_faithful_sentence(candidate, abstract)

def build_bibliography_entry(paper: dict, style: str, number: int) -> str:
    authors = get_all_author_names(paper)
    title = paper.get("title") or "Untitled"
    venue = paper.get("venue") or "Unknown venue"
    year = paper.get("year") or "n.d."
    url = paper.get("url") or ""
    
    if style == "IEEE":
        return f'[{number}] {authors}, "{title}," *{venue}*, {year}. Available: {url}'
    elif style == "APA":
        return f"{authors} ({year}). {title}. *{venue}*. {url}"
    else:  # Harvard
        return f"{authors} ({year}) '{title}', *{venue}*. Available at: {url}"

def get_next_ieee_number(bibliography: dict) -> int:
    if not bibliography:
        return 1
    return max(entry["number"] for entry in bibliography.values()) + 1

def register_bibliography_entry(paper: dict, style: str, bibliography: dict) -> int:
    key = paper.get("url") or paper.get("title", "unknown")
    if key in bibliography:
        return bibliography[key]["number"]
    number = get_next_ieee_number(bibliography)
    bibliography[key] = {"number": number, "paper": paper}
    return number

# ---------------------------------------------------------------------------
# PROOFREADING (OpenAI) + DIFF
# ---------------------------------------------------------------------------

def proofread_text(text: str) -> str:
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        prompt = (
            "You are an elite academic editor and proofreader. Your task is to proofread and elevate the provided text.\n\n"
            "### OBJECTIVES:\n"
            "- Correct all spelling, grammar, punctuation, and syntax errors.\n"
            "- Enhance clarity, flow, and conciseness while maintaining a rigorous academic tone.\n"
            "- Implement all improvements directly into the text (do not just suggest them).\n\n"
            "### STRICT CONSTRAINTS:\n"
            "1. NO COMMENTARY: Return strictly the revised text and nothing else. Do not include introductory phrases, greetings, or explanations.\n"
            "2. PRESERVE FORMATTING: You must strictly maintain all original Markdown formatting (headings like #, bold like **, italics like *, lists, and paragraph breaks).\n"
            "3. PRESERVE MEANING: Do not alter the original data, facts, or core arguments.\n\n"
            f"### ORIGINAL TEXT:\n{text}"
        )
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional academic editor."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        revised = response.choices[0].message.content.strip()
        return revised
    except Exception as e:
        app.logger.error(f"Proofreading error: {e}")
        return None

def generate_diff_html(original: str, revised: str) -> str:
    original_lines = original.splitlines()
    revised_lines = revised.splitlines()
    diff = difflib.HtmlDiff(wrapcolumn=80)
    html = diff.make_table(original_lines, revised_lines, context=True, numlines=3)
    style = """
    <style>
        .diff { font-family: monospace; font-size: 0.8rem; line-height: 1.4; }
        .diff td { padding: 1px 4px; }
        .diff_header { background: #f0ebe1; }
        .diff_add { background: #d4edda; color: #155724; }
        .diff_sub { background: #f8d7da; color: #721c24; }
        .diff_chg { background: #fff3cd; color: #856404; }
        .diff_ctx { background: #fcf9f5; }
    </style>
    """
    return style + html

# ---------------------------------------------------------------------------
# SESSION INIT
# ---------------------------------------------------------------------------

def ensure_session():
    if 'editor_text' not in session:
        session['editor_text'] = DEFAULT_DOCUMENT_TEXT
    if 'bibliography' not in session:
        session['bibliography'] = {}
    if 'citation_style' not in session:
        session['citation_style'] = 'IEEE'
    if 'fetched_papers' not in session:
        session['fetched_papers'] = []
    if 'fetch_error' not in session:
        session['fetch_error'] = None
    if 'active_query' not in session:
        session['active_query'] = ''
    if 'mode' not in session:
        session['mode'] = 'citation'

# ---------------------------------------------------------------------------
# FLASK ROUTES
# ---------------------------------------------------------------------------

@app.route('/')
def home():
    ensure_session()
    return render_template('home.html')

@app.route('/citation')
def citation_mode():
    ensure_session()
    session['mode'] = 'citation'
    text = session['editor_text']
    paragraphs = split_into_paragraphs(text)
    citable_indices = [i for i, p in enumerate(paragraphs) if is_citable_paragraph(p)]
    bibliography = session['bibliography']
    style = session['citation_style']
    bib_list = []
    for entry in bibliography.values():
        ref = build_bibliography_entry(entry['paper'], style, entry['number'])
        bib_list.append({'number': entry['number'], 'reference': ref})
    bib_list.sort(key=lambda x: x['number'])
    return render_template(
        'citation.html',
        editor_text=text,
        citation_style=style,
        bibliography=bib_list,
        fetched_papers=session.get('fetched_papers', []),
        fetch_error=session['fetch_error'],
        active_query=session['active_query'],
        paragraphs=paragraphs,
        citable_indices=citable_indices
    )

@app.route('/proofreading')
def proofreading_mode():
    ensure_session()
    session['mode'] = 'proofreading'
    text = session['editor_text']
    return render_template('proofreading.html', editor_text=text)

# ---- API routes for Citation mode ----

@app.route('/set_style', methods=['POST'])
def set_style():
    style = request.form.get('citation_style')
    if style in ('IEEE', 'APA', 'Harvard'):
        session['citation_style'] = style
    return redirect(url_for('citation_mode'))

@app.route('/clear_bibliography', methods=['POST'])
def clear_bibliography():
    session['bibliography'] = {}
    return redirect(url_for('citation_mode'))

@app.route('/fetch_paper', methods=['POST'])
def fetch_paper():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Invalid request'}), 400
    editor_text = data.get('editor_text')
    paragraph_index = data.get('paragraph_index')
    if editor_text is None or paragraph_index is None:
        return jsonify({'error': 'Missing editor_text or paragraph_index'}), 400
    session['editor_text'] = editor_text
    paragraphs = split_into_paragraphs(editor_text)
    if paragraph_index >= len(paragraphs):
        return jsonify({'error': 'Paragraph index out of range'}), 400
    paragraph = paragraphs[paragraph_index]
    if not is_citable_paragraph(paragraph):
        return jsonify({'error': 'Selected paragraph is not citable (too short or heading).'}), 400

    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        prompt = (
            f"Extract exactly {KEYWORDS_PER_QUERY} highly relevant academic search keywords "
            f"or short compound phrases from the following text. "
            f"Return ONLY the keywords separated by commas, nothing else.\n\n"
            f"Text: {paragraph}"
        )
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert academic librarian."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2
        )
        raw_result = response.choices[0].message.content
        keywords = [kw.strip() for kw in raw_result.split(',')]
        if not keywords:
            session['fetch_error'] = 'Could not extract meaningful keywords from this paragraph.'
            session['fetched_papers'] = []
            session['active_query'] = ''
            return jsonify({'error': session['fetch_error']}), 200
    except Exception as e:
        app.logger.error(f"OpenAI error: {e}")
        session['fetch_error'] = 'OpenAI error. Please check your API key.'
        session['fetched_papers'] = []
        session['active_query'] = ''
        return jsonify({'error': session['fetch_error']}), 200

    search_query = ' '.join(keywords)
    session['active_query'] = search_query
    papers, error = fetch_top_papers(search_query)
    if error:
        session['fetch_error'] = error
        session['fetched_papers'] = []
        return jsonify({'error': error}), 200
        
    session['fetched_papers'] = papers
    session['fetch_error'] = None
    return jsonify({
        'papers': papers,
        'query': search_query
    })

@app.route('/insert_citation', methods=['POST'])
def insert_citation():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Invalid request'}), 400
    editor_text = data.get('editor_text')
    paragraph_index = data.get('paragraph_index')
    insertion_mode = data.get('insertion_mode')
    paper_index = data.get('paper_index', 0)
    
    if not all([editor_text, paragraph_index is not None, insertion_mode]):
        return jsonify({'error': 'Missing required fields'}), 400

    papers = session.get('fetched_papers', [])
    if not papers:
        active_query = session.get('active_query', '')
        if active_query in _cached_papers:
            papers = _cached_papers[active_query]['papers']

    if not papers or paper_index >= len(papers):
        return jsonify({'error': 'Invalid paper selected. Please fetch papers again.'}), 400

    paper = papers[paper_index]
    style = session['citation_style']
    bibliography = session['bibliography']

    number = register_bibliography_entry(paper, style, bibliography)
    session['bibliography'] = bibliography

    paragraphs = split_into_paragraphs(editor_text)
    if paragraph_index >= len(paragraphs):
        return jsonify({'error': 'Paragraph index out of range'}), 400

    current = paragraphs[paragraph_index].rstrip()
    tag = build_inline_tag(paper, style, number)
    
    # ---- EXTRACTION DE CITATION INTELLIGENTE AVEC OPENAI ----
    inline = f" {tag}" # Format par défaut
    
    if insertion_mode == "Insert Quote + Citation":
        name_year_tag = build_name_year_tag(paper)  # toujours (Nom, Année) DANS LE TEXTE
        source_text = get_citable_source_text(paper)  # abstract, sinon scraping de la page du papier
        if source_text:
            try:
                quote = extract_faithful_quote(current, source_text)
                inline = f' "{quote}" {name_year_tag}'
            except Exception as e:
                app.logger.error(f"OpenAI quote extraction error: {e}")
                inline = f" {name_year_tag}"
        else:
            # S'il n'y a pas de résumé pour cet article dans la base de données, 
            # on insère juste la balise de citation.
            inline = f" {name_year_tag}"

    # Gestion propre de la ponctuation (Le point s'insère APRES la citation)
    if current and current[-1] in ('.', ','):
        punct = current[-1]
        base = current[:-1].rstrip()
        paragraphs[paragraph_index] = f"{base} {inline.strip()}{punct}"
    else:
        paragraphs[paragraph_index] = f"{current} {inline.strip()}"

    new_text = "\n\n".join(paragraphs)
    session['editor_text'] = new_text

    bib_list = []
    for entry in bibliography.values():
        ref = build_bibliography_entry(entry['paper'], style, entry['number'])
        bib_list.append({'number': entry['number'], 'reference': ref})
    bib_list.sort(key=lambda x: x['number'])

    return jsonify({
        'editor_text': new_text,
        'bibliography': bib_list
    })

# ---- Upload ----

@app.route('/upload_document', methods=['POST'])
def upload_document():
    if 'document' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['document']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    try:
        file_bytes = file.read()
        text = extract_text_from_upload(file_bytes, file.filename)
        session['editor_text'] = text
        session['fetched_papers'] = []
        session['fetch_error'] = None
        session['active_query'] = ''
        paragraphs = split_into_paragraphs(text)
        citable_indices = [i for i, p in enumerate(paragraphs) if is_citable_paragraph(p)]
        return jsonify({
            'editor_text': text,
            'paragraphs': paragraphs,
            'citable_indices': citable_indices
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 400

# ---- Download DOCX ----

@app.route('/download_document', methods=['GET', 'POST'])
def download_document():
    ensure_session()
    
    if request.method == 'POST':
        text = request.form.get('current_text', '')
    else:
        text = session.get('editor_text', '')
        
    mode = session.get('mode', 'citation')
    
    doc = docx.Document()
    
    for p in text.split('\n\n'):
        add_markdown_paragraph_to_docx(doc, p)

    if mode == 'citation':
        bibliography = session.get('bibliography', {})
        if bibliography:
            doc.add_page_break()
            doc.add_heading('References', level=1)
            style = session['citation_style']
            sorted_entries = sorted(bibliography.values(), key=lambda x: x['number'])
            for entry in sorted_entries:
                ref = build_bibliography_entry(entry['paper'], style, entry['number'])
                doc.add_paragraph(ref, style='List Bullet')
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(
        output,
        as_attachment=True,
        download_name='document.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# ---- Proofreading API ----

@app.route('/proofread', methods=['POST'])
def proofread():
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided'}), 400
        text = data['text']
        revised = proofread_text(text)
        if revised is None:
            return jsonify({'error': 'Proofreading failed. Please check your OpenAI API key.'}), 500
        diff_html = generate_diff_html(text, revised)
        return jsonify({
            'revised_text': revised,
            'diff_html': diff_html,
            'original_text': text
        })
    except Exception as e:
        app.logger.error(f"Proofreading route error: {e}")
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/update_text', methods=['POST'])
def update_text():
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400
    session['editor_text'] = data['text']
    return jsonify({'success': True})

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)